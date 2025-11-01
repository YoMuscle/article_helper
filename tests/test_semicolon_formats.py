"""
測試分號分隔引用的各種格式變化
"""
import sys
from pathlib import Path

# 添加專案根目錄到 sys.path
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from services.document_analyzer import DocumentAnalyzer
from docx import Document
from docx.shared import Pt
import tempfile
import os

def create_test_document(text):
    """創建一個測試用的 Word 文檔"""
    doc = Document()
    
    # 添加標題
    doc.add_heading('Test Document', 0)
    
    # 添加內文段落
    para = doc.add_paragraph(text)
    
    # 添加參考文獻區段
    doc.add_heading('References', 1)
    doc.add_paragraph('Wang, L., & Smith, J. (2015). First study. Journal A, 10(1), 1-10.')
    doc.add_paragraph('Cooke, M. (2015). Second study. Journal B, 15(2), 20-30.')
    doc.add_paragraph('Lee, K., et al. (2016). Third study. Journal C, 20(3), 40-50.')
    
    # 保存到臨時文件
    fd, temp_path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    doc.save(temp_path)
    return temp_path

def test_semicolon_variations():
    """測試各種分號分隔引用的格式變化"""
    
    test_cases = [
        {
            'name': '正確格式：分號後有空格',
            'text': 'Previous studies (Wang & Smith, 2015; Cooke, 2015) have shown this.',
            'should_have_errors': False
        },
        {
            'name': '錯誤格式：分號後缺少空格',
            'text': 'Previous studies (Wang & Smith, 2015;Cooke, 2015) have shown this.',
            'should_have_errors': True
        },
        {
            'name': '錯誤格式：分號前多了空格',
            'text': 'Previous studies (Wang & Smith, 2015 ; Cooke, 2015) have shown this.',
            'should_have_errors': True
        },
        {
            'name': '錯誤格式：分號前後都有多餘空格',
            'text': 'Previous studies (Wang & Smith, 2015 ;  Cooke, 2015) have shown this.',
            'should_have_errors': True
        },
        {
            'name': '缺少左括號（malformed）',
            'text': 'Previous studies Wang & Smith, 2015; Cooke, 2015) have shown this.',
            'should_have_errors': True
        },
        {
            'name': '三個引用：正確格式',
            'text': 'Studies (Wang & Smith, 2015; Cooke, 2015; Lee et al., 2016) support this.',
            'should_have_errors': False
        },
        {
            'name': '三個引用：第二個引用缺少空格',
            'text': 'Studies (Wang & Smith, 2015;Cooke, 2015; Lee et al., 2016) support this.',
            'should_have_errors': True
        }
    ]
    
    print("\n" + "="*80)
    print("測試分號分隔引用的格式檢查")
    print("="*80)
    
    analyzer = DocumentAnalyzer()
    passed = 0
    failed = 0
    
    for i, test_case in enumerate(test_cases, 1):
        print(f"\n測試 {i}: {test_case['name']}")
        print(f"測試文本: {test_case['text']}")
        
        # 創建測試文檔
        temp_path = create_test_document(test_case['text'])
        
        try:
            # 分析文檔
            results = analyzer.analyze_document(temp_path)
            
            # 檢查格式錯誤
            has_format_errors = False
            if 'format_errors' in results and results['format_errors']:
                has_format_errors = True
                print(f"  發現格式錯誤: {len(results['format_errors'])} 個")
                for error in results['format_errors']:
                    # 適應不同的錯誤格式
                    if isinstance(error, dict):
                        citation = error.get('citation', error.get('text', 'Unknown'))
                        error_msg = error.get('error', error.get('message', ''))
                        error_msgs = error.get('errors', [])
                        if error_msg:
                            print(f"    - {citation}: {error_msg}")
                        elif error_msgs:
                            if isinstance(error_msgs, str):
                                error_msgs = [error_msgs]
                            print(f"    - {citation}: {', '.join(error_msgs)}")
                        else:
                            print(f"    - {citation}: (無錯誤訊息)")
                            print(f"      完整錯誤物件: {error}")
                    else:
                        print(f"    - {error}")
            
            # 檢查是否符合預期
            expected = test_case['should_have_errors']
            actual = has_format_errors
            
            if expected == actual:
                print(f"  ✓ 通過（預期{'有' if expected else '無'}錯誤，實際{'有' if actual else '無'}錯誤）")
                passed += 1
            else:
                print(f"  ✗ 失敗（預期{'有' if expected else '無'}錯誤，但實際{'有' if actual else '無'}錯誤）")
                failed += 1
                
                # 顯示詳細資訊以便除錯
                if 'citations' in results:
                    print(f"  找到的引用:")
                    for cit in results['citations']:
                        print(f"    - {cit['text']} (類型: {cit.get('type', 'unknown')})")
        finally:
            # 刪除臨時文件
            if os.path.exists(temp_path):
                os.unlink(temp_path)
    
    print("\n" + "="*80)
    print(f"測試結果: {passed} 通過, {failed} 失敗")
    print("="*80)
    
    return failed == 0

if __name__ == '__main__':
    success = test_semicolon_variations()
    sys.exit(0 if success else 1)
