"""
測試 Lopez-Calderon 複合姓氏的匹配
"""
import sys
from pathlib import Path

project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from services.document_analyzer import DocumentAnalyzer
from docx import Document
import tempfile
import os

def create_test_document():
    doc = Document()
    doc.add_heading('Test Document', 0)
    
    # 測試文本：引用 Lopez-Calderon
    text = 'The EEG data were preprocessed using EEGLAB (Lopez-Calderon & Luck, 2014) functions.'
    doc.add_paragraph(text)
    
    # 參考文獻
    doc.add_heading('References', 1)
    doc.add_paragraph('Lopez-Calderon, J., & Luck, S. J. (2014). ERPLAB: an open-source toolbox for the analysis of event-related potentials. Frontiers in Human Neuroscience, 8, 213.')
    
    fd, temp_path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    doc.save(temp_path)
    return temp_path

def test_lopez_calderon():
    print("\n" + "="*80)
    print("測試 Lopez-Calderon 複合姓氏的匹配")
    print("="*80)
    
    analyzer = DocumentAnalyzer()
    temp_path = create_test_document()
    
    try:
        results = analyzer.analyze_document(temp_path)
        
        print(f"\n【參考文獻數量】: {results.get('total_references', 0)}")
        print(f"【引用數量】: {results.get('total_citations', 0)}")
        
        print("\n【參考文獻狀態】")
        for ref in results.get('citation_status', []):
            status = "✓ 已引用" if ref.get('cited') else "✗ 未引用"
            print(f"  {status}: {ref.get('reference', 'Unknown')[:80]}...")
        
        print("\n【格式錯誤】")
        if results.get('format_errors'):
            for error in results['format_errors']:
                print(f"  ✗ {error['citation']}")
                print(f"    錯誤: {error['error']}")
        else:
            print("  無格式錯誤")
        
        print("\n【缺失/錯誤的引用】")
        if results.get('missing_references'):
            for missing in results['missing_references']:
                print(f"  ✗ {missing['citation']}")
                if missing.get('suggestion'):
                    print(f"    💡 建議: {missing['suggestion']}")
        else:
            print("  無缺失的引用")
        
        print("\n" + "="*80)
        if results.get('citation_status'):
            is_cited = results['citation_status'][0].get('cited', False)
            if is_cited:
                print("✅ 成功！Lopez-Calderon & Luck (2014) 被正確識別為已引用")
            else:
                print("❌ 失敗！Lopez-Calderon & Luck (2014) 未被識別為已引用")
        print("="*80)
        
    finally:
        if os.path.exists(temp_path):
            os.unlink(temp_path)

if __name__ == '__main__':
    test_lopez_calderon()
