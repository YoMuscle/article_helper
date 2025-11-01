"""
測試第 511 行的格式錯誤
"""
import sys
from pathlib import Path

# 添加專案根目錄到 sys.path
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from services.document_analyzer import DocumentAnalyzer
from docx import Document
import tempfile
import os

def create_test_document(text):
    """創建一個測試用的 Word 文檔"""
    doc = Document()
    
    # 添加標題
    doc.add_heading('Test Document', 0)
    
    # 添加內文段落
    para = doc.add_paragraph(text)
    
    # 添加參考文獻區段
    doc.add_heading('References', 1)
    doc.add_paragraph('Klimesch, W. (1999). EEG alpha and theta oscillations reflect cognitive and memory performance: A review.')
    doc.add_paragraph('Klimesch, W., Sauseng, P., & Hanslmayr, S. (2007). EEG alpha oscillations: The inhibition-timing hypothesis.')
    doc.add_paragraph('Cooke, M. (2015). Test article. Journal, 10, 1-10.')
    doc.add_paragraph('Lopez-Calderon, J., & Luck, S. J. (2014). ERPLAB: An open-source toolbox for the analysis of event-related potentials.')
    
    # 保存到臨時文件
    fd, temp_path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    doc.save(temp_path)
    return temp_path

def test_format_error():
    """測試格式錯誤問題"""
    
    print("\n" + "="*80)
    print("測試格式錯誤問題")
    print("="*80)
    
    # 測試案例 1: 三個引用，格式看起來正確
    test_text = 'executive processes operating on task-relevant information that is stored in working memory (Klimesch 1999; Klimesch and Sauseng, 2007; Cooke, 2015).'
    
    print(f"\n測試文本: {test_text}")
    print(f"預期: 應該檢測到格式錯誤（缺少逗號）")
    
    analyzer = DocumentAnalyzer()
    temp_path = create_test_document(test_text)
    
    try:
        results = analyzer.analyze_document(temp_path)
        
        # 顯示找到的引用
        print(f"\n找到的引用: {len(results.get('citations', []))} 個")
        for i, cit in enumerate(results.get('citations', []), 1):
            print(f"  {i}. {cit['text']}")
            print(f"     類型: {cit.get('type', 'unknown')}")
            print(f"     位置: {cit.get('position', 'unknown')}")
            if cit.get('from_multi_citation'):
                print(f"     來自多重引用: True")
            if cit.get('original_text'):
                print(f"     原始文本: {cit.get('original_text')}")
        
        # 顯示格式錯誤
        if 'format_errors' in results and results['format_errors']:
            print(f"\n格式錯誤: {len(results['format_errors'])} 個")
            for error in results['format_errors']:
                print(f"  - {error['citation']}")
                print(f"    錯誤: {error['error']}")
        else:
            print("\n✓ 沒有格式錯誤")
        
        # 顯示缺失的參考文獻
        if 'missing_references' in results and results['missing_references']:
            print(f"\n缺失的參考文獻: {len(results['missing_references'])} 個")
            for missing in results['missing_references']:
                print(f"  - {missing['citation']}")
                if missing.get('suggestion'):
                    print(f"    建議: {missing['suggestion']}")
        else:
            print("\n✓ 沒有缺失的參考文獻")
        
        # 顯示參考文獻狀態
        print(f"\n參考文獻狀態:")
        for ref_id, ref_data in results.get('references', {}).items():
            status = "✓ 已引用" if ref_data.get('cited') else "✗ 未引用"
            print(f"  {status}: {ref_data.get('parenthetical', 'Unknown')}")
    
    finally:
        if os.path.exists(temp_path):
            os.unlink(temp_path)

if __name__ == '__main__':
    test_format_error()
