"""
測試三位作者但引用時只寫兩位的情況
"""
import sys
from pathlib import Path

project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from services.document_analyzer import DocumentAnalyzer
from docx import Document
import tempfile
import os

def create_test_document():
    doc = Document()
    doc.add_heading('Test Document', 0)
    
    # 測試文本：引用有問題
    text = 'executive processes operating on task-relevant information (Klimesch, 1999; Klimesch and Sauseng, 2007; Cooke, 2015).'
    doc.add_paragraph(text)
    
    # 參考文獻：Klimesch 2007 有 3 位作者
    doc.add_heading('References', 1)
    doc.add_paragraph('Klimesch, W. (1999). EEG alpha and theta oscillations reflect cognitive and memory performance: A review.')
    doc.add_paragraph('Klimesch, W., Sauseng, P., & Hanslmayr, S. (2007). EEG alpha oscillations: The inhibition-timing hypothesis.')
    doc.add_paragraph('Cooke, M. (2015). Test article. Journal, 10, 1-10.')
    
    fd, temp_path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    doc.save(temp_path)
    return temp_path

def test_three_authors_cited_as_two():
    print("\n" + "="*80)
    print("測試：3 位作者的參考文獻被引用為 2 位作者")
    print("="*80)
    
    analyzer = DocumentAnalyzer()
    temp_path = create_test_document()
    
    try:
        results = analyzer.analyze_document(temp_path)
        
        print(f"\n【調試信息】")
        print(f"  結果keys: {list(results.keys())}")
        print(f"  參考文獻數量: {results.get('total_references', 0)}")
        print(f"  引用數量: {results.get('total_citations', 0)}")
        
        print("\n【參考文獻狀態】")
        for ref in results.get('citation_status', []):
            print(f"  - {ref.get('reference', 'Unknown')}")
            print(f"    已引用: {'✓' if ref.get('cited') else '✗'}")
        
        print("\n【格式錯誤】")
        if results.get('format_errors'):
            for error in results['format_errors']:
                print(f"  ✗ {error['citation']}")
                print(f"    錯誤: {error['error']}")
        else:
            print("  無格式錯誤")
        
        print("\n【缺失/錯誤的引用】")
        if results.get('missing_references'):
            for missing in results['missing_references']:
                print(f"  ✗ {missing['citation']}")
                print(f"    章節: {missing.get('section', 'Unknown')}")
                if missing.get('suggestion'):
                    print(f"    💡 建議: {missing['suggestion']}")
        else:
            print("  無缺失的引用")
        
        print("\n" + "="*80)
        print("分析")
        print("="*80)
        print("問題: 文中寫 'Klimesch and Sauseng, 2007'")
        print("但是: Klimesch et al. (2007) 有 3 位作者")
        print("應該: 引用應該寫成 'Klimesch et al., 2007'")
        
    finally:
        if os.path.exists(temp_path):
            os.unlink(temp_path)

if __name__ == '__main__':
    test_three_authors_cited_as_two()
